"""Create minimal Word (.docx) files for case compose flows."""

from __future__ import annotations

import re
from datetime import date
from pathlib import Path
from typing import Any

# ---------------------------------------------------------------------------
# Precedent merge codes
# ---------------------------------------------------------------------------

def _precedent_code_suffix(slot: int) -> str:
    """Merge key suffix for additional clients 2–4, e.g. [TITLE] -> [TITLE_2]."""
    return f"_{slot}]"


def _merge_key_with_suffix(code: str, slot: int) -> str:
    if not code.startswith("[") or not code.endswith("]"):
        return code
    return code[:-1] + _precedent_code_suffix(slot)


# Name / company codes that are repeated for additional clients 2, 3 & 4 (see build_merge_fields).
_ADDITIONAL_CLIENT_NAME_CODES: tuple[str, ...] = (
    "[TITLE]",
    "[FIRST_NAME]",
    "[FIRST_INITIAL]",
    "[MIDDLE_NAME]",
    "[MIDDLE_INITIAL]",
    "[LAST_NAME]",
    "[LAST_INITIAL]",
    "[COMPANY_NAME]",
    "[TRADING_NAME]",
)

# Lawyer matter contacts are organisation-only; merge codes for the lawyer row are company / trading only.
_LAWYER_ROW_NAME_CODES: tuple[str, ...] = ("[COMPANY_NAME]", "[TRADING_NAME]")


PRECEDENT_CODES: dict[str, str] = {
    # Person
    "[TITLE]": "Title (e.g. Mr / Mrs / Dr)",
    "[FIRST_NAME]": "First name",
    "[FIRST_INITIAL]": "First initial (e.g. J)",
    "[MIDDLE_NAME]": "Middle name",
    "[MIDDLE_INITIAL]": "Middle initial",
    "[LAST_NAME]": "Surname",
    "[LAST_INITIAL]": "Surname initial",
    # Organisation
    "[COMPANY_NAME]": "Registered company name",
    "[TRADING_NAME]": "Trading name",
    # Address (shared)
    "[ADDR1]": "Address line 1",
    "[ADDR2]": "Address line 2",
    "[ADDR3]": "Town / city",
    "[ADDR4]": "County",
    "[POSTCODE]": "Postcode",
    # Case / matter
    "[MATTER_DESCRIPTION]": "Matter description",
    "[CASE_REF]": "Case reference number",
    "[DATE]": "Date when the document is generated (DD/MM/YYYY)",
    "[FEE_EARNER]": "Fee earner display name (from the case fee earner)",
    "[FEE_EARNER_JOB_TITLE]": "Fee earner job title (from the case fee earner user)",
    "[CONTACT_REF]": "Contact's reference (as stored in canary)",
}

for _slot_num, _slot_label in ((2, "2nd"), (3, "3rd"), (4, "4th")):
    for _base_key in _ADDITIONAL_CLIENT_NAME_CODES:
        _suff_key = _merge_key_with_suffix(_base_key, _slot_num)
        PRECEDENT_CODES[_suff_key] = (
            f"{PRECEDENT_CODES[_base_key]} — additional client {_slot_num} "
            f"({_slot_label} 'Client' matter contact on the case, by date added)"
        )

for _li in range(1, 5):
    for _base_key in _LAWYER_ROW_NAME_CODES:
        _inner = _base_key[1:-1]
        _lk = f"[LAWYER_{_li}_{_inner}]"
        PRECEDENT_CODES[_lk] = (
            f"Lawyer {_li}: {PRECEDENT_CODES[_base_key]} "
            "(among 'Lawyers' matter contacts, by date added; lawyers are organisation contacts)"
        )

for _li in range(1, 5):
    for _cj in range(1, 5):
        for _base_key in _ADDITIONAL_CLIENT_NAME_CODES:
            _inner = _base_key[1:-1]
            _lk = f"[LAWYER_{_li}_CLIENT_{_cj}_{_inner}]"
            PRECEDENT_CODES[_lk] = f"Lawyer {_li}'s linked client {_cj}: {PRECEDENT_CODES[_base_key]}"

# Extra fields on each lawyer-linked client (name/company codes are in the loop above).
_LAWYER_LINKED_CLIENT_EXTRA: tuple[tuple[str, str], ...] = (
    ("NAME", "Display name on the contact card"),
    ("TYPE", "person or organisation"),
    ("EMAIL", "Email"),
    ("PHONE", "Phone"),
    ("ADDR1", "Address line 1"),
    ("ADDR2", "Address line 2"),
    ("ADDR3", "Town / city"),
    ("ADDR4", "County"),
    ("POSTCODE", "Postcode"),
    ("COUNTRY", "Country"),
    ("MATTER_REFERENCE", "Matter-specific reference on this case"),
    ("MATTER_CONTACT_TYPE", "Matter contact type label on this case"),
)

for _li in range(1, 5):
    for _cj in range(1, 5):
        for _inner, _lab in _LAWYER_LINKED_CLIENT_EXTRA:
            _lk = f"[LAWYER_{_li}_CLIENT_{_cj}_{_inner}]"
            PRECEDENT_CODES[_lk] = f"Lawyer {_li}'s linked client {_cj}: {_lab} (Case matter contact)"

# Shorthand: same values as [LAWYER_1_CLIENT_cj_*] (first Lawyers matter contact on the case, by date added).
_LAWYER_CONTACT_CLIENT_ALIAS_INNERS: tuple[str, ...] = tuple(
    c[1:-1] for c in _ADDITIONAL_CLIENT_NAME_CODES
) + tuple(x[0] for x in _LAWYER_LINKED_CLIENT_EXTRA)

for _cj in range(1, 5):
    for _inner in _LAWYER_CONTACT_CLIENT_ALIAS_INNERS:
        _lk = f"[LAWYER_CONTACT_CLIENT_{_cj}_{_inner}]"
        PRECEDENT_CODES[_lk] = (
            f"Same as [LAWYER_1_CLIENT_{_cj}_{_inner}]: first 'Lawyers' matter contact’s linked client {_cj} "
            "(by date added among Lawyers contacts)"
        )

# Explicit “selected in compose” contact (letter/document precedent); filled when a contact is chosen in the UI.
_CONTACT_COMPOSE_STATIC: tuple[tuple[str, str], ...] = (
    ("[CONTACT_NAME]", "Display name on the contact card"),
    ("[CONTACT_TYPE]", "person or organisation"),
    ("[CONTACT_EMAIL]", "Email"),
    ("[CONTACT_PHONE]", "Phone"),
    ("[CONTACT_ADDR1]", "Address line 1"),
    ("[CONTACT_ADDR2]", "Address line 2"),
    ("[CONTACT_ADDR3]", "Town / city"),
    ("[CONTACT_ADDR4]", "County"),
    ("[CONTACT_POSTCODE]", "Postcode"),
    ("[CONTACT_COUNTRY]", "Country"),
    (
        "[CONTACT_MATTER_REFERENCE]",
        "Matter-specific reference (case contact snapshot only; empty for a global directory contact)",
    ),
    (
        "[CONTACT_MATTER_CONTACT_TYPE]",
        "Matter contact type label on this case (case contact only; empty for a global directory contact)",
    ),
)
for _ck, _desc in _CONTACT_COMPOSE_STATIC:
    PRECEDENT_CODES[_ck] = (
        f"Selected contact for this compose: {_desc}. Empty if no contact was chosen in the dialogue."
    )

for _base_key in _ADDITIONAL_CLIENT_NAME_CODES:
    _inner = _base_key[1:-1]
    _cc_key = f"[CONTACT_{_inner}]"
    PRECEDENT_CODES[_cc_key] = (
        f"Selected contact for this compose: same field as [{_inner}] for person or organisation name parts; "
        f"always the contact picked in the dialogue (including when “merge all clients” fills [{_inner}] from another client)."
    )


def _s_str(v: object) -> str:
    return (v or "").strip() if isinstance(v, str) else ""


def _initial_letter(v: object) -> str:
    t = _s_str(v)
    return t[0].upper() if t else ""


def _core_name_company_for_contact(contact: Any | None) -> dict[str, str]:
    """Nine merge keys shared by primary and additional-client slots."""

    if not contact:
        return {k: "" for k in _ADDITIONAL_CLIENT_NAME_CODES}

    contact_type = _s_str(getattr(contact, "type", "person"))
    first = _s_str(getattr(contact, "first_name", None))
    middle = _s_str(getattr(contact, "middle_name", None))
    last = _s_str(getattr(contact, "last_name", None))

    company = _s_str(getattr(contact, "company_name", None))
    if not company and contact_type == "organisation":
        company = _s_str(getattr(contact, "name", None))

    return {
        "[TITLE]": _s_str(getattr(contact, "title", None)),
        "[FIRST_NAME]": first,
        "[FIRST_INITIAL]": _initial_letter(first),
        "[MIDDLE_NAME]": middle,
        "[MIDDLE_INITIAL]": _initial_letter(middle),
        "[LAST_NAME]": last,
        "[LAST_INITIAL]": _initial_letter(last),
        "[COMPANY_NAME]": company,
        "[TRADING_NAME]": _s_str(getattr(contact, "trading_name", None)),
    }


def _contact_type_str(contact: Any) -> str:
    t = getattr(contact, "type", None)
    if t is None:
        return ""
    if hasattr(t, "value"):
        return str(t.value)
    return str(t)


def _lawyer_linked_client_extra_map(contact: Any | None) -> dict[str, str]:
    """Extra merge fields for a lawyer-linked CaseContact (beyond the nine name/company keys)."""

    if not contact:
        return {x[0]: "" for x in _LAWYER_LINKED_CLIENT_EXTRA}
    return {
        "NAME": _s_str(getattr(contact, "name", None)),
        "TYPE": _contact_type_str(contact),
        "EMAIL": _s_str(getattr(contact, "email", None)),
        "PHONE": _s_str(getattr(contact, "phone", None)),
        "ADDR1": _s_str(getattr(contact, "address_line1", None)),
        "ADDR2": _s_str(getattr(contact, "address_line2", None)),
        "ADDR3": _s_str(getattr(contact, "city", None)),
        "ADDR4": _s_str(getattr(contact, "county", None)),
        "POSTCODE": _s_str(getattr(contact, "postcode", None)),
        "COUNTRY": _s_str(getattr(contact, "country", None)),
        "MATTER_REFERENCE": _s_str(getattr(contact, "matter_contact_reference", None)),
        "MATTER_CONTACT_TYPE": _s_str(getattr(contact, "matter_contact_type", None)),
    }


def _empty_precedent_field_map() -> dict[str, str]:
    return {k: "" for k in PRECEDENT_CODES}


def _apply_lawyer_merge_slots(
    out: dict[str, str],
    lawyer_slots: list[tuple[Any, list[Any]] | None] | None,
) -> None:
    if not lawyer_slots:
        return
    for i in range(min(4, len(lawyer_slots))):
        slot = lawyer_slots[i]
        if not slot:
            continue
        law_cc, client_list = slot
        li = i + 1
        law_core = _core_name_company_for_contact(law_cc)
        for merge_key in _LAWYER_ROW_NAME_CODES:
            inner = merge_key[1:-1]
            out[f"[LAWYER_{li}_{inner}]"] = law_core.get(merge_key, "")
        for j, cli in enumerate((client_list or [])[:4]):
            cj = j + 1
            ccore = _core_name_company_for_contact(cli)
            for merge_key, val in ccore.items():
                inner = merge_key[1:-1]
                out[f"[LAWYER_{li}_CLIENT_{cj}_{inner}]"] = val
            extras = _lawyer_linked_client_extra_map(cli)
            for inner, val in extras.items():
                out[f"[LAWYER_{li}_CLIENT_{cj}_{inner}]"] = val
            if li == 1:
                for merge_key, val in ccore.items():
                    inner = merge_key[1:-1]
                    out[f"[LAWYER_CONTACT_CLIENT_{cj}_{inner}]"] = val
                for inner, val in extras.items():
                    out[f"[LAWYER_CONTACT_CLIENT_{cj}_{inner}]"] = val


def _fill_compose_selected_contact_codes(out: dict[str, str], contact: Any | None) -> None:
    """Fill ``[CONTACT_*]`` keys from the contact chosen in the compose dialogue (if any)."""

    if not contact:
        return
    core = _core_name_company_for_contact(contact)
    for merge_key in _ADDITIONAL_CLIENT_NAME_CODES:
        inner = merge_key[1:-1]
        out[f"[CONTACT_{inner}]"] = core.get(merge_key, "")
    out["[CONTACT_NAME]"] = _s_str(getattr(contact, "name", None))
    out["[CONTACT_TYPE]"] = _contact_type_str(contact)
    out["[CONTACT_EMAIL]"] = _s_str(getattr(contact, "email", None))
    out["[CONTACT_PHONE]"] = _s_str(getattr(contact, "phone", None))
    out["[CONTACT_ADDR1]"] = _s_str(getattr(contact, "address_line1", None))
    out["[CONTACT_ADDR2]"] = _s_str(getattr(contact, "address_line2", None))
    out["[CONTACT_ADDR3]"] = _s_str(getattr(contact, "city", None))
    out["[CONTACT_ADDR4]"] = _s_str(getattr(contact, "county", None))
    out["[CONTACT_POSTCODE]"] = _s_str(getattr(contact, "postcode", None))
    out["[CONTACT_COUNTRY]"] = _s_str(getattr(contact, "country", None))
    out["[CONTACT_MATTER_REFERENCE]"] = _s_str(getattr(contact, "matter_contact_reference", None))
    out["[CONTACT_MATTER_CONTACT_TYPE]"] = _s_str(getattr(contact, "matter_contact_type", None))


def build_merge_fields(
    case: Any,
    fee_earner_name: str = "",
    fee_earner_job_title: str = "",
    merge_date: date | None = None,
    *,
    merge_all_clients: bool = False,
    ordered_client_contacts: list[Any] | None = None,
    selected_contact: Any | None = None,
    selected_client_slot: int | None = None,
    lawyer_slots: list[tuple[Any, list[Any]] | None] | None = None,
    compose_selected_contact: Any | None = None,
) -> dict[str, str]:
    """Build precedent code→value dict.

    * **merge_all_clients** — Fill client 1 from ``ordered_client_contacts[0]`` into unsuffixed
      keys and ``[ADDR*]``; clients 2–4 into ``[TITLE_2]`` … ``[TRADING_NAME_4]``.
      ``[CONTACT_REF]`` is taken from the first client row.

    * **Single Client matter contact** (``selected_client_slot`` 1–4) — Fill only that client’s
      name/company keys (slot 1 unsuffixed; slots 2–4 use ``_2`` … ``_4``). Address and
      ``[CONTACT_REF]`` come from ``selected_contact``.

    * **Global contact or non-Client matter contact** (``selected_client_slot`` is None) — Fill
      unsuffixed name and address keys only; suffixed client keys stay empty.

    * **compose_selected_contact** — When set (the contact chosen in the compose UI), fills
      ``[CONTACT_*]`` codes from that row even when ``merge_all_clients`` is True, so templates
      can address the picked contact separately from unsuffixed client merge keys.
    """

    out = _empty_precedent_field_map()

    def finalize(m: dict[str, str]) -> dict[str, str]:
        _apply_lawyer_merge_slots(m, lawyer_slots)
        _fill_compose_selected_contact_codes(m, compose_selected_contact)
        return m

    matter_desc = _s_str(getattr(case, "title", None)) if case else ""
    case_ref = _s_str(getattr(case, "case_number", None)) if case else ""
    d = merge_date or date.today()
    date_str = d.strftime("%d/%m/%Y")

    out["[MATTER_DESCRIPTION]"] = matter_desc
    out["[CASE_REF]"] = case_ref
    out["[DATE]"] = date_str
    out["[FEE_EARNER]"] = fee_earner_name
    out["[FEE_EARNER_JOB_TITLE]"] = fee_earner_job_title

    oc = [c for c in (ordered_client_contacts or [])][:4]

    if merge_all_clients:
        for i, cc in enumerate(oc):
            core = _core_name_company_for_contact(cc)
            if i == 0:
                for k, v in core.items():
                    out[k] = v
                out["[ADDR1]"] = _s_str(getattr(cc, "address_line1", None))
                out["[ADDR2]"] = _s_str(getattr(cc, "address_line2", None))
                out["[ADDR3]"] = _s_str(getattr(cc, "city", None))
                out["[ADDR4]"] = _s_str(getattr(cc, "county", None))
                out["[POSTCODE]"] = _s_str(getattr(cc, "postcode", None))
                out["[CONTACT_REF]"] = _s_str(getattr(cc, "matter_contact_reference", None))
            else:
                slot = i + 1
                for k, v in core.items():
                    out[_merge_key_with_suffix(k, slot)] = v
        return finalize(out)

    if selected_contact is None:
        return finalize(out)

    contact_ref = _s_str(getattr(selected_contact, "matter_contact_reference", None))
    out["[ADDR1]"] = _s_str(getattr(selected_contact, "address_line1", None))
    out["[ADDR2]"] = _s_str(getattr(selected_contact, "address_line2", None))
    out["[ADDR3]"] = _s_str(getattr(selected_contact, "city", None))
    out["[ADDR4]"] = _s_str(getattr(selected_contact, "county", None))
    out["[POSTCODE]"] = _s_str(getattr(selected_contact, "postcode", None))

    if selected_client_slot is None or not (1 <= selected_client_slot <= 4):
        core = _core_name_company_for_contact(selected_contact)
        for k, v in core.items():
            out[k] = v
        out["[CONTACT_REF]"] = contact_ref
        return finalize(out)

    idx = selected_client_slot - 1
    cc = oc[idx] if idx < len(oc) else None
    if cc is None:
        out["[CONTACT_REF]"] = contact_ref
        return finalize(out)

    core = _core_name_company_for_contact(cc)
    if selected_client_slot == 1:
        for k, v in core.items():
            out[k] = v
    else:
        for k, v in core.items():
            out[_merge_key_with_suffix(k, selected_client_slot)] = v

    out["[CONTACT_REF]"] = contact_ref
    return finalize(out)


def _replace_in_text(text: str, fields: dict[str, str]) -> str:
    # Longest keys first so a shorter placeholder can never break a longer token (defensive).
    for code in sorted(fields.keys(), key=len, reverse=True):
        text = text.replace(code, fields[code])
    return text


# Inner token without brackets, e.g. TITLE, LAST_NAME_3 — for slot detection.
_NAME_CODE_INNERS: frozenset[str] = frozenset(c[1:-1] for c in _ADDITIONAL_CLIENT_NAME_CODES)

_PLACEHOLDER_TOKEN_RE = re.compile(r"\[([A-Z0-9_]+)\]")


def _name_slot_from_placeholder_inner(inner: str) -> int | None:
    """Return 1–4 for per-client name/company placeholders; ``None`` for other codes."""
    if inner in _NAME_CODE_INNERS:
        return 1
    m = re.fullmatch(r"(.+)_([234])$", inner)
    if not m:
        return None
    base, suf = m.group(1), m.group(2)
    if base not in _NAME_CODE_INNERS:
        return None
    return int(suf)


def _contact_has_any_name_or_company_field(contact: Any | None) -> bool:
    core = _core_name_company_for_contact(contact)
    return any((v or "").strip() for v in core.values())


def _inter_client_sep_flags(ordered_clients: list[Any] | None) -> dict[tuple[int, int], bool]:
    """When True, insert `` and `` between adjacent name placeholders for slots (a, b)."""
    if not ordered_clients:
        return {}
    out: dict[tuple[int, int], bool] = {}
    n = min(len(ordered_clients), 4)
    for i in range(n - 1):
        a, b = ordered_clients[i], ordered_clients[i + 1]
        if _contact_has_any_name_or_company_field(a) and _contact_has_any_name_or_company_field(b):
            out[(i + 1, i + 2)] = True
    return out


def _insert_and_between_adjacent_name_placeholders(
    text: str,
    sep_flags: dict[tuple[int, int], bool],
) -> str:
    """Insert the word ``and`` between consecutive client name placeholders when ``sep_flags`` says to.

    Matches only **adjacent** ``[CODE]`` tokens in this string (same paragraph). Whitespace
    between them is replaced by `` and `` (spaces around *and*).
    """
    if not sep_flags:
        return text
    matches = list(_PLACEHOLDER_TOKEN_RE.finditer(text))
    if len(matches) < 2:
        return text
    parts: list[str] = []
    pos = 0
    i = 0
    while i < len(matches):
        m = matches[i]
        parts.append(text[pos : m.start()])
        parts.append(m.group(0))
        pos = m.end()
        if i + 1 < len(matches):
            m2 = matches[i + 1]
            s1 = _name_slot_from_placeholder_inner(m.group(1))
            s2 = _name_slot_from_placeholder_inner(m2.group(1))
            if (
                s1 is not None
                and s2 is not None
                and s2 == s1 + 1
                and sep_flags.get((s1, s2), False)
            ):
                parts.append(" and ")
                pos = m2.start()
        i += 1
    parts.append(text[pos:])
    return "".join(parts)


def _xml_escape_ooxml_text(value: str) -> str:
    """Escape text merged into ``<w:t>`` (and similar) XML character data."""
    from xml.sax.saxutils import escape

    return escape(value, {'"': "&quot;", "'": "&apos;"})


def _ooxml_part_paths_for_merge() -> tuple[str, ...]:
    """Part paths inside the .docx zip that may contain visible merge tokens."""
    return (
        "word/document.xml",
        "word/footnotes.xml",
        "word/endnotes.xml",
    )


def _ooxml_part_path_matches(name: str) -> bool:
    if name in _ooxml_part_paths_for_merge():
        return True
    if name.startswith("word/header") and name.endswith(".xml"):
        return True
    if name.startswith("word/footer") and name.endswith(".xml"):
        return True
    return False


def _merge_precedent_codes_in_ooxml_zip(src_bytes: bytes, fields: dict[str, str]) -> bytes:
    """Replace ``[CODE]`` substrings in raw OOXML parts (one pass per file).

    Runs **after** the python-docx paragraph pass. Catches any remaining contiguous
    placeholders in XML (including footnotes) and tokens split across ``<w:t>`` boundaries
    that the paragraph walk could not join for ``and`` insertion.
    """
    import io
    import zipfile

    escaped = {k: _xml_escape_ooxml_text(v) for k, v in fields.items()}
    src = io.BytesIO(src_bytes)
    out = io.BytesIO()
    with zipfile.ZipFile(src, "r") as zin, zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for info in zin.infolist():
            raw = zin.read(info.filename)
            if _ooxml_part_path_matches(info.filename):
                try:
                    text = raw.decode("utf-8")
                except UnicodeDecodeError:
                    zout.writestr(info, raw)
                    continue
                text = _replace_in_text(text, escaped)
                raw = text.encode("utf-8")
            zout.writestr(info, raw)
    return out.getvalue()


def _rewrite_paragraph_to_single_run(para: Any, replaced: str) -> None:
    """Replace paragraph content with plain text in one run.

    Word often puts merge tokens in hyperlinked or oddly split runs. Clearing only ``para.runs``
    can leave ``w:hyperlink`` / nested ``w:t`` behind, so the old token still appears next to
    the merged text (e.g. surname twice with a gap). We strip non-``w:pPr`` children and add
    a single run — same approach as a clean retype of the paragraph.
    """
    from docx.oxml.ns import qn

    p_el = para._p
    for child in list(p_el):
        if child.tag != qn("w:pPr"):
            p_el.remove(child)
    para.add_run(replaced)


def merge_precedent_codes(
    src_bytes: bytes,
    fields: dict[str, str],
    *,
    ordered_clients: list[Any] | None = None,
    merge_all_clients: bool = False,
) -> bytes:
    """Replace [CODE] placeholders in a .docx (precedent merge).

    1. **python-docx paragraph pass** — when ``merge_all_clients`` is true, inserts the word
       ``and`` between adjacent client name placeholders for consecutive clients that both
       have name/company data; then substitutes fields; handles merged table cells; removes
       code-only blank paragraphs.

    2. **Zip / OOXML pass** — replaces any remaining contiguous ``[CODE]`` substrings in
       document parts (including split tokens not fixed in step 1).
    """
    sep_flags = _inter_client_sep_flags(ordered_clients) if merge_all_clients else {}
    merged = _merge_precedent_codes_via_python_docx(src_bytes, fields, sep_flags)
    return _merge_precedent_codes_in_ooxml_zip(merged, fields)


def _merge_precedent_codes_via_python_docx(
    src_bytes: bytes,
    fields: dict[str, str],
    sep_flags: dict[tuple[int, int], bool],
) -> bytes:
    """Paragraph walk: optional *and* insertion, field replace, blank-line cleanup."""
    import io
    from docx import Document

    doc = Document(io.BytesIO(src_bytes))

    # Build a regex that matches any known code so we can detect code-only paragraphs
    code_pattern = re.compile("|".join(re.escape(k) for k in fields))
    seen_wp: set[Any] = set()

    def _merge_para(para: Any) -> bool:
        """Merge codes in para. Returns True if the para should be removed (became blank)."""
        wp = para._p
        if wp in seen_wp:
            return False
        # Use Paragraph.text, not join(para.runs): runs inside w:hyperlink are not top-level runs,
        # and leaving those w:t nodes caused duplicate visible text after merging the first run.
        full = para.text
        if not full:
            return False  # already empty — don't touch
        if sep_flags:
            full = _insert_and_between_adjacent_name_placeholders(full, sep_flags)
        had_code = bool(code_pattern.search(full))
        if not had_code:
            return False
        # Claim only once we will rewrite, so empty / no-code paragraphs visited from duplicate
        # merged cells can still be processed on a later distinct visit (should not happen, but safe).
        seen_wp.add(wp)
        replaced = _replace_in_text(full, fields)
        _rewrite_paragraph_to_single_run(para, replaced)
        # Remove the paragraph if it's now blank (was code-only, value was empty)
        return not replaced.strip()

    def _process_paras(paras: Any) -> None:
        plist = list(paras) if not isinstance(paras, list) else paras
        to_remove = [p for p in plist if _merge_para(p)]
        for p in to_remove:
            p._element.getparent().remove(p._element)

    def _iter_distinct_cells(table: Any):
        """Each physical cell once (merged cells share one ``w:tc`` but span multiple grid slots)."""
        seen_tc: set[Any] = set()
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                if tc in seen_tc:
                    continue
                seen_tc.add(tc)
                yield cell

    def _process_table(table: Any) -> None:
        for cell in _iter_distinct_cells(table):
            _process_paras(cell.paragraphs)
            nested = getattr(cell, "tables", None)
            if nested:
                for nt in nested:
                    _process_table(nt)

    # Body paragraphs (not inside tables)
    _process_paras(doc.paragraphs)

    for table in doc.tables:
        _process_table(table)

    # Headers / footers
    for section in doc.sections:
        for hf in (section.header, section.footer,
                   section.even_page_header, section.even_page_footer,
                   section.first_page_header, section.first_page_footer):
            if hf.is_linked_to_previous:
                continue
            _process_paras(hf.paragraphs)
            for table in hf.tables:
                _process_table(table)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _set_default_proofing_language_en_gb(doc: Any) -> None:
    """Set OOXML default run language to en-GB for new documents.

    ONLYOFFICE/Word use this for the default *document* proofing language (spell check).
    Editor JWT ``lang``/``region`` do not replace language embedded in existing .docx files.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    styles_el = doc.styles.element
    dd = styles_el.find(qn("w:docDefaults"))
    if dd is None:
        return
    rpd = dd.find(qn("w:rPrDefault"))
    if rpd is None:
        return
    rpr = rpd.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        rpd.append(rpr)
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), "en-GB")
    lang.set(qn("w:eastAsia"), "en-GB")
    lang.set(qn("w:bidi"), "en-GB")


def write_blank_docx(path: Path) -> None:
    from docx import Document

    doc = Document()
    _set_default_proofing_language_en_gb(doc)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def write_completion_statement_docx(
    path: Path,
    *,
    case_number: str,
    client_name: str | None,
    finance: Any,  # FinanceOut (or dict with .categories list)
) -> None:
    """Write a completion statement .docx from case finance data."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _fmt_pence(p: int | None) -> str:
        if p is None:
            return ""
        val = abs(p) / 100
        return f"\u00a3{val:,.2f}"  # £ with thousands separator

    def _set_cell_shading(cell, fill: str) -> None:
        """Apply a background fill colour (hex) to a table cell."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        tcPr.append(shd)

    def _set_cell_borders(cell, top=None, bottom=None, left=None, right=None) -> None:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
            if val:
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:val"), val.get("val", "single"))
                el.set(qn("w:sz"), str(val.get("sz", 4)))
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), val.get("color", "auto"))
                tcBorders.append(el)
        tcPr.append(tcBorders)

    doc = Document()
    _set_default_proofing_language_en_gb(doc)

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ── Title ─────────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("COMPLETION STATEMENT")
    run.bold = True
    run.font.size = Pt(16)

    # ── Sub-header: matter reference + date ───────────────────────────────────
    matter_line = case_number
    if client_name:
        matter_line = f"{case_number} — {client_name}"
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(matter_line)
    sub_run.font.size = Pt(11)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Date: {date.today().strftime('%d %B %Y')}")
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()  # spacer

    # ── Main table ────────────────────────────────────────────────────────────
    # Columns: Description | Debit | Credit
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, label in enumerate(("Description", "Debit", "Credit")):
        cell = hdr_cells[i]
        cell.text = label
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
        _set_cell_shading(cell, "D0DAEA")

    # Column widths (Description wide, Debit/Credit equal)
    col_widths = [Inches(3.8), Inches(1.5), Inches(1.5)]
    for i, width in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = width

    total_dr = 0
    total_cr = 0

    categories = getattr(finance, "categories", None) or []

    for cat in categories:
        cat_name = getattr(cat, "name", None) or str(cat)
        items = getattr(cat, "items", None) or []

        # Category header row
        row = table.add_row()
        row.cells[0].merge(row.cells[2])
        merged = row.cells[0]
        merged.text = cat_name.upper()
        run = merged.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)
        _set_cell_shading(merged, "EEF2F8")

        for item in items:
            name = getattr(item, "name", "") or ""
            direction = getattr(item, "direction", "debit")
            amount_pence = getattr(item, "amount_pence", None)

            if direction == "debit":
                debit_str = _fmt_pence(amount_pence)
                credit_str = ""
                if amount_pence:
                    total_dr += amount_pence
            else:
                debit_str = ""
                credit_str = _fmt_pence(amount_pence)
                if amount_pence:
                    total_cr += amount_pence

            row = table.add_row()
            row.cells[0].text = name
            row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
            row.cells[1].text = debit_str
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row.cells[1].paragraphs[0].runs[0 if row.cells[1].paragraphs[0].runs else -1].font.size = Pt(10) if row.cells[1].paragraphs[0].runs else None
            row.cells[2].text = credit_str
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if row.cells[2].paragraphs[0].runs:
                row.cells[2].paragraphs[0].runs[0].font.size = Pt(10)
            for ci in range(3):
                if row.cells[ci].paragraphs[0].runs:
                    row.cells[ci].paragraphs[0].runs[0].font.size = Pt(10)

    # ── Totals row ────────────────────────────────────────────────────────────
    tot_row = table.add_row()
    tot_row.cells[0].text = "TOTALS"
    tot_row.cells[1].text = _fmt_pence(total_dr)
    tot_row.cells[2].text = _fmt_pence(total_cr)
    for ci, cell in enumerate(tot_row.cells):
        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(cell.text)
        run.bold = True
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
        _set_cell_shading(cell, "D0DAEA")

    # ── Balance row ───────────────────────────────────────────────────────────
    balance = total_cr - total_dr
    bal_row = table.add_row()
    bal_row.cells[0].merge(bal_row.cells[1])
    bal_label = bal_row.cells[0]
    bal_label.text = "BALANCE DUE FROM CLIENT" if balance > 0 else "BALANCE DUE TO CLIENT" if balance < 0 else "BALANCE"
    bal_run = bal_label.paragraphs[0].runs[0] if bal_label.paragraphs[0].runs else bal_label.paragraphs[0].add_run(bal_label.text)
    bal_run.bold = True
    bal_run.font.size = Pt(10)
    _set_cell_shading(bal_label, "EEF2F8")

    bal_val_cell = bal_row.cells[2]
    bal_val_cell.text = _fmt_pence(abs(balance))
    bal_val_run = bal_val_cell.paragraphs[0].runs[0] if bal_val_cell.paragraphs[0].runs else bal_val_cell.paragraphs[0].add_run(bal_val_cell.text)
    bal_val_run.bold = True
    bal_val_run.font.size = Pt(10)
    bal_val_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_cell_shading(bal_val_cell, "EEF2F8")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def extract_plain_text_from_docx_bytes(data: bytes) -> str:
    """Best-effort plain text from a .docx for e-mail body (M365 Graph)."""
    from io import BytesIO

    from docx import Document

    doc = Document(BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = (p.text or "").strip()
                    if t:
                        parts.append(t)
    return "\n\n".join(parts) if parts else ""
